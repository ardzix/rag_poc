"""
Service untuk ekstraksi teks dari berbagai format dokumen
"""
import re
from typing import Tuple, Optional, Dict, Any
from datetime import date, datetime
import PyPDF2
import docx
import magic
import openpyxl
from openpyxl.utils import get_column_letter


class DocumentExtractor:
    """
    Service untuk mengekstrak teks dari dokumen PDF, DOCX, dan TXT
    """
    
    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
    }
    
    @staticmethod
    def detect_mime_type(file_obj) -> str:
        """
        Deteksi MIME type dari file object
        """
        try:
            # Reset file pointer
            file_obj.seek(0)
            
            # Baca header untuk deteksi
            mime = magic.from_buffer(file_obj.read(2048), mime=True)
            
            # Reset lagi
            file_obj.seek(0)
            
            # Fallback untuk XLSX yang kadang terdeteksi sebagai application/zip
            filename = getattr(file_obj, 'name', '') or ''
            if mime == 'application/zip' and filename.lower().endswith('.xlsx'):
                return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            
            return mime
        except Exception:
            # Fallback ke content_type dari upload
            return file_obj.content_type if hasattr(file_obj, 'content_type') else 'application/octet-stream'
    
    @staticmethod
    def is_supported(mime_type: str) -> bool:
        """
        Cek apakah MIME type didukung
        """
        return mime_type in DocumentExtractor.SUPPORTED_MIME_TYPES
    
    @staticmethod
    def extract(file_obj, mime_type: str) -> Tuple[str, Optional[str], Optional[Dict[str, Any]]]:
        """
        Ekstrak teks dari file
        
        Args:
            file_obj: File object dari upload
            mime_type: MIME type yang sudah terdeteksi
        
        Returns:
            Tuple (extracted_text, error_message, structured_data)
            Jika berhasil: (text, None, structured_data or None)
            Jika gagal: ("", error_message, None)
        """
        try:
            file_obj.seek(0)
            
            if mime_type == 'application/pdf':
                text, err = DocumentExtractor._extract_pdf(file_obj)
                return (text, err, None)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                text, err = DocumentExtractor._extract_docx(file_obj)
                return (text, err, None)
            elif mime_type == 'text/plain':
                text, err = DocumentExtractor._extract_txt(file_obj)
                return (text, err, None)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                return DocumentExtractor._extract_xlsx(file_obj)
            else:
                return ("", f"MIME type tidak didukung: {mime_type}", None)
                
        except Exception as e:
            return ("", f"Error saat ekstraksi: {str(e)}", None)
    
    @staticmethod
    def _extract_pdf(file_obj) -> Tuple[str, Optional[str]]:
        """Ekstrak teks dari PDF"""
        try:
            reader = PyPDF2.PdfReader(file_obj)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            if not text_parts:
                return ("", "PDF tidak mengandung teks yang bisa diekstrak (mungkin hasil scan)")
            
            combined_text = "\n".join(text_parts)
            normalized_text = DocumentExtractor._normalize_text(combined_text)
            
            return (normalized_text, None)
            
        except Exception as e:
            return ("", f"Gagal membaca PDF: {str(e)}")
    
    @staticmethod
    def _extract_docx(file_obj) -> Tuple[str, Optional[str]]:
        """Ekstrak teks dari DOCX"""
        try:
            doc = docx.Document(file_obj)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Ekstrak tabel juga
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            if not text_parts:
                return ("", "Dokumen DOCX kosong")
            
            combined_text = "\n".join(text_parts)
            normalized_text = DocumentExtractor._normalize_text(combined_text)
            
            return (normalized_text, None)
            
        except Exception as e:
            return ("", f"Gagal membaca DOCX: {str(e)}")
    
    @staticmethod
    def _extract_txt(file_obj) -> Tuple[str, Optional[str]]:
        """Ekstrak teks dari TXT"""
        try:
            # Coba decode dengan beberapa encoding
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file_obj.seek(0)
                    text = file_obj.read().decode(encoding)
                    normalized_text = DocumentExtractor._normalize_text(text)
                    return (normalized_text, None)
                except UnicodeDecodeError:
                    continue
            
            return ("", "Tidak bisa decode file TXT dengan encoding yang didukung")
            
        except Exception as e:
            return ("", f"Gagal membaca TXT: {str(e)}")
    
    @staticmethod
    def _normalize_cell_value(value):
        """
        Normalisasi nilai cell agar JSON-serializable dan konsisten
        """
        if isinstance(value, datetime):
            return value.isoformat()
        if isinstance(value, date):
            return value.isoformat()
        return value
    
    @staticmethod
    def _extract_xlsx(file_obj) -> Tuple[str, Optional[str], Optional[Dict[str, Any]]]:
        """Ekstrak teks dan struktur dari XLSX"""
        try:
            wb = openpyxl.load_workbook(file_obj, data_only=True)
            sheets_data = []
            summary_parts = []
            
            for ws in wb.worksheets:
                # Ambil semua rows
                raw_rows = list(ws.iter_rows(values_only=True))
                
                if not raw_rows:
                    continue
                
                # Tentukan header
                header_row = raw_rows[0]
                if any(cell is not None and str(cell).strip() for cell in header_row):
                    columns = [str(cell).strip() if cell is not None else "" for cell in header_row]
                    data_rows = raw_rows[1:]
                else:
                    # Fallback: gunakan A, B, C...
                    max_cols = max(len(r) for r in raw_rows)
                    columns = [get_column_letter(i + 1) for i in range(max_cols)]
                    data_rows = raw_rows
                
                # Normalisasi rows
                rows = []
                for row in data_rows:
                    normalized_row = [
                        DocumentExtractor._normalize_cell_value(cell)
                        for cell in row
                    ]
                    rows.append(normalized_row)
                
                sheet_info = {
                    "name": ws.title,
                    "columns": columns,
                    "rows": rows
                }
                sheets_data.append(sheet_info)
                
                # Summary text per sheet (untuk konteks LLM)
                row_count = len(rows)
                preview_rows = rows[:5]
                summary_parts.append(
                    f"Sheet: {ws.title}\n"
                    f"Kolom: {', '.join(columns)}\n"
                    f"Total Rows: {row_count}\n"
                    f"Contoh Rows (maks 5): {preview_rows}\n"
                )
            
            if not sheets_data:
                return ("", "Dokumen XLSX kosong", None)
            
            summary_text = "\n".join(summary_parts)
            normalized_text = DocumentExtractor._normalize_text(summary_text)
            
            structured_data = {
                "format": "xlsx",
                "sheets": sheets_data
            }
            
            return (normalized_text, None, structured_data)
            
        except Exception as e:
            return ("", f"Gagal membaca XLSX: {str(e)}", None)
    
    @staticmethod
    def _normalize_text(text: str) -> str:
        """
        Normalisasi teks:
        - Hilangkan karakter non-printable
        - Rapikan whitespace
        - Hilangkan baris kosong berlebihan
        """
        # Hilangkan karakter kontrol kecuali newline dan tab
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalisasi whitespace: ganti multiple spaces dengan single space
        text = re.sub(r' +', ' ', text)
        
        # Normalisasi newline: maksimal 2 newline berturut-turut
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Trim setiap baris
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
